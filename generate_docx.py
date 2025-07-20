from docx import Document
from datetime import datetime
import io
import json
from docx.shared import Pt
from copy import deepcopy
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_row_with_style(table, template_row, row_data, improvement_lookup):
    new_row = table.add_row()
    for i, cell in enumerate(template_row.cells):
        # 스타일 복사
        new_row.cells[i]._tc.get_or_add_tcPr().append(deepcopy(cell._tc.get_or_add_tcPr()))
    
    equipment = row_data.get("equipment", "")
    risk = row_data.get("risk", "")
    evaluation = row_data.get("evaluation", "")
    
    new_row.cells[0].text = equipment
    new_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_row.cells[1].text = risk

    # 적정
    if evaluation == "적정":
        new_row.cells[2].text = "◯"
        new_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 보완
    elif evaluation == "보완":
        new_row.cells[3].text = "◯"
        new_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_row.cells[4].text = improvement_lookup.get((equipment, risk), "")
    
    # 나머지 셀은 공란 (개선완료일, 담당자)
    return new_row

def generate_docx_file(data):
    doc = Document("base/위험성평가_견본파일.docx")
    table = doc.tables[0]

    # 개선사항 사전 만들기
    with open("data/위평_데이터_변환.json", "r", encoding="utf-8") as f:
        raw_json = json.load(f)

    improvement_lookup = {}
    for entry in raw_json:
        equipment = entry["설비"]
        for item in entry["유해위험요인목록"]:
            key = (equipment, item["유해위험요인"])
            improvement_lookup[key] = item["개선사항"]

    # 템플릿 행 복사용 (3번째 행)
    template_row = table.rows[2]

    # 기존 공란 행 채우기
    filled_rows = 0
    for row in table.rows[3:]:
        if all(cell.text.strip() == "" for cell in row.cells):
            if filled_rows < len(data):
                # 기존 공란 행 삭제 후 새 행 삽입
                table._tbl.remove(row._tr)
                insert_row_with_style(table, template_row, data[filled_rows], improvement_lookup)
                filled_rows += 1
            else:
                break

    # 남은 데이터는 새 행 추가
    for row_data in data[filled_rows:]:
        insert_row_with_style(table, template_row, row_data, improvement_lookup)

    # 제목 날짜 치환
    for paragraph in doc.paragraphs:
        if "날짜" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(datetime.today().strftime("%Y-%m-%d") + " 위험성 평가")
            run.font.size = Pt(14)
            break

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
